"""Unit tests for modules.quiz_exporter.word_renderer."""
from __future__ import annotations

import re
import tempfile
from pathlib import Path

import pytest
from docx.oxml.ns import qn

from modules.quiz_exporter.word_renderer import (
    ExamMeta,
    ExportConfig,
    ExportQuestionSnapshot,
    PrintProfile,
    WordRenderer,
    build_output_path,
)
from core.utils.latex_rendering import render_inline_latex_text

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

MC_Q = {
    "type": "MC",
    "content": "Câu hỏi MC?",
    "point_value": 1.0,
    "difficulty": "easy",
    "learning_outcome_code": "CLO_1",
    "category": "Chương 1",
    "hint": "",
    "explanation": "",
    "options": [
        {"key": "A", "text": "Đáp án A", "is_correct": True},
        {"key": "B", "text": "Đáp án B", "is_correct": False},
        {"key": "C", "text": "Đáp án C", "is_correct": False},
    ],
}

MA_Q = {
    "type": "MA",
    "content": "Câu hỏi MA?",
    "point_value": 2.0,
    "difficulty": "medium",
    "learning_outcome_code": "CLO_2",
    "category": "Chương 1",
    "hint": "",
    "explanation": "",
    "options": [
        {"key": "A", "text": "Đáp án A", "is_correct": True},
        {"key": "B", "text": "Đáp án B", "is_correct": True},
        {"key": "C", "text": "Đáp án C", "is_correct": False},
    ],
}

BLANK_Q = {
    "type": "BLANK",
    "content": "Điền vào chỗ trống: ___ là thủ đô của Việt Nam.",
    "point_value": 1.5,
    "difficulty": "easy",
    "learning_outcome_code": "CLO_2",
    "category": "Chương 2",
    "hint": "",
    "explanation": "",
    "accepted_answers": ["Hà Nội", "Ha Noi"],
}

SA_Q = {
    "type": "SA",
    "content": "Giải thích khái niệm trọng lực.",
    "point_value": 3.0,
    "difficulty": "hard",
    "learning_outcome_code": "CLO_3",
    "category": "Chương 3",
    "hint": "",
    "explanation": "",
    "accepted_answers": ["lực hút của Trái Đất"],
}

PROBLEM_Q = {
    "type": "ES",
    "content": "Giải bài toán kiểm định giả thuyết.",
    "point_value": 6.0,
    "difficulty": "Phân tích",
    "learning_outcome_code": "CLO_4",
    "category": "Chương 4",
    "hint": "",
    "explanation": "",
    "accepted_answers": ["Nêu giả thuyết", "Tính thống kê kiểm định"],
    "question_variant": "problem",
    "problem_rubric": [
        {"marker": "B1", "content": "Nêu giả thuyết", "score": 2.0},
        {"marker": "B2", "content": "Tính thống kê kiểm định", "score": 4.0},
    ],
}

PROBLEM_Q_TEMPLATE = {
    **PROBLEM_Q,
    "problem_template_id": 33,
    "problem_template_name": "Mẫu kiểm định chuẩn",
}

PROBLEM_Q_LONG = {
    "type": "ES",
    "content": "Giải bài toán dài.",
    "point_value": 8.0,
    "difficulty": "Đánh giá",
    "learning_outcome_code": "CLO_5",
    "category": "Chương 5",
    "hint": "",
    "explanation": "",
    "accepted_answers": [
        "Dòng đáp án rất dài thứ nhất",
        "Dòng đáp án rất dài thứ hai",
    ],
    "question_variant": "problem",
    "problem_rubric": [
        {
            "marker": "B1",
            "content": "Bước 1: Phân tích đề bài rất dài để đảm bảo nội dung đầy đủ trong file đáp án được xuất ra.",
            "score": 3.0,
        },
        {
            "marker": "B1",
            "content": "Bước 2: Xuống dòng kiểm tra\nVà vẫn phải hiện đủ nội dung ở ô đáp án mẫu.",
            "score": 2.0,
        },
        {
            "marker": "B2",
            "content": "Bước 3: Kết luận cuối cùng và nêu rõ ý nghĩa thống kê của kết quả.",
            "score": 3.0,
        },
    ],
}

PROBLEM_Q_LATEX = {
    "type": "ES",
    "content": "Kiểm tra công thức $t=\\frac{\\bar{x}-\\mu_0}{s/\\sqrt{n}}$ trong bài toán.",
    "point_value": 5.0,
    "difficulty": "Phân tích",
    "learning_outcome_code": "CLO_6",
    "category": "Chương 6",
    "hint": "",
    "explanation": "",
    "accepted_answers": ["$t=\\frac{\\bar{x}-\\mu_0}{s/\\sqrt{n}}$"],
    "question_variant": "problem",
    "problem_rubric": [
        {
            "marker": "B1",
            "content": "Viết $t=\\frac{\\bar{x}-\\mu_0}{s/\\sqrt{n}}$ và kiểm tra điều kiện.",
            "score": 5.0,
        },
    ],
}

TF_Q = {
    "type": "TF",
    "content": "Trái Đất quay quanh Mặt Trời.",
    "point_value": 1.0,
    "difficulty": "hard",
    "learning_outcome_code": "CLO_1",
    "category": "Chương 2",
    "hint": "",
    "explanation": "",
    "options": [
        {"key": "A", "text": "Đúng", "is_correct": True},
        {"key": "B", "text": "Sai", "is_correct": False},
    ],
}

ALL_QUESTIONS = [MC_Q, MA_Q, TF_Q, BLANK_Q, SA_Q]


def _text(doc) -> str:
    """Concatenate all paragraph and table-cell text in a Document."""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
    return "\n".join(parts)


def _cell_width_twips(cell) -> int:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    assert tc_w is not None
    return int(tc_w.get(qn("w:w")))


# ---------------------------------------------------------------------------
# ExamMeta defaults
# ---------------------------------------------------------------------------

class TestExamMeta:
    def test_defaults(self):
        meta = ExamMeta()
        assert meta.exam_title == "BÀI KIỂM TRA"
        assert meta.school == ""
        assert meta.duration_minutes == 0

    def test_custom(self):
        meta = ExamMeta(school="Trường XYZ", exam_title="Giữa kỳ", duration_minutes=60)
        assert meta.school == "Trường XYZ"
        assert meta.exam_title == "Giữa kỳ"
        assert meta.duration_minutes == 60


# ---------------------------------------------------------------------------
# ExportConfig defaults
# ---------------------------------------------------------------------------

class TestExportConfig:
    def test_defaults(self):
        cfg = ExportConfig()
        assert cfg.show_instructions is True
        assert cfg.show_answer_sheet is True
        assert cfg.show_answer_key is True
        assert cfg.show_scoring_rules is True
        assert cfg.show_question_points is False
        assert cfg.show_question_statistics is False
        assert cfg.numbering_mode == "global"
        assert cfg.group_by_type is True
        assert cfg.show_cover_sheet is False
        assert cfg.split_answer_key_file is False
        assert cfg.watermark_text == ""
        assert cfg.cover_sheet_template == "standard"
        assert cfg.answer_key_naming_policy == "suffix"


class TestPrintProfile:
    def test_defaults(self):
        profile = PrintProfile()
        assert profile.page_size == "A4"
        assert profile.show_student_info_block is True
        assert profile.left_margin_cm == pytest.approx(2.0)


# ---------------------------------------------------------------------------
# _types_present
# ---------------------------------------------------------------------------

class TestTypesPresent:
    def _renderer(self):
        return WordRenderer()

    def test_all_types(self):
        r = self._renderer()
        assert r._types_present(ALL_QUESTIONS) == {"MC", "MA", "TF", "BLANK", "SA"}

    def test_single_type(self):
        r = self._renderer()
        assert r._types_present([MC_Q, MC_Q]) == {"MC"}

    def test_empty(self):
        r = self._renderer()
        assert r._types_present([]) == set()


# ---------------------------------------------------------------------------
# _group_questions
# ---------------------------------------------------------------------------

class TestGroupQuestions:
    def _renderer(self):
        return WordRenderer()

    def test_flat_mode(self):
        r = self._renderer()
        cfg = ExportConfig(group_by_type=False)
        groups = r._group_questions(ALL_QUESTIONS, cfg)
        assert len(groups) == 1
        letter, title, qs = groups[0]
        assert letter == ""
        assert title == ""
        assert len(qs) == 5

    def test_grouped_mode_all_types(self):
        r = self._renderer()
        cfg = ExportConfig(group_by_type=True)
        groups = r._group_questions(ALL_QUESTIONS, cfg)
        # MC→A, MA→B, TF→C, BLANK→D, SA→E
        assert len(groups) == 5
        letters = [g[0] for g in groups]
        assert letters == ["A", "B", "C", "D", "E"]

    def test_grouped_mode_mc_only(self):
        r = self._renderer()
        cfg = ExportConfig(group_by_type=True)
        groups = r._group_questions([MC_Q, MC_Q], cfg)
        assert len(groups) == 1
        letter, title, qs = groups[0]
        assert letter == "A"
        assert "Multiple Choice" in title
        assert len(qs) == 2

    def test_grouped_mode_mc_and_blank(self):
        r = self._renderer()
        cfg = ExportConfig(group_by_type=True)
        groups = r._group_questions([MC_Q, BLANK_Q], cfg)
        assert len(groups) == 2
        # MC is section A, BLANK is section B
        assert groups[0][0] == "A"
        assert groups[1][0] == "B"
        assert "Điền vào chỗ trống" in groups[1][1]

    def test_preserves_order_within_group(self):
        r = self._renderer()
        cfg = ExportConfig(group_by_type=True)
        q1 = {**MC_Q, "content": "First MC"}
        q2 = {**MC_Q, "content": "Second MC"}
        groups = r._group_questions([q1, q2], cfg)
        _, _, qs = groups[0]
        assert qs[0]["content"] == "First MC"
        assert qs[1]["content"] == "Second MC"


# ---------------------------------------------------------------------------
# render() — basic structure
# ---------------------------------------------------------------------------

class TestRender:
    def _render(self, questions=None, meta=None, config=None):
        if questions is None:
            questions = ALL_QUESTIONS
        if meta is None:
            meta = ExamMeta(exam_title="Test Exam")
        if config is None:
            config = ExportConfig()
        return WordRenderer().render(questions, meta, config)

    def test_returns_document(self):
        doc = self._render()
        # docx.Document is a factory function; check via the concrete type
        import docx.document
        assert isinstance(doc, docx.document.Document)

    def test_accepts_typed_export_snapshots(self):
        typed_questions = [ExportQuestionSnapshot.from_dict(MC_Q)]
        doc = self._render(questions=typed_questions)
        text = _text(doc)
        assert "Câu hỏi MC?" in text

    def test_exam_title_in_output(self):
        meta = ExamMeta(exam_title="BÀI KIỂM TRA CUỐI KỲ")
        doc = self._render(meta=meta)
        text = _text(doc)
        assert "BÀI KIỂM TRA CUỐI KỲ" in text

    def test_school_in_output(self):
        meta = ExamMeta(exam_title="T", school="Trường ABC")
        doc = self._render(meta=meta)
        text = _text(doc)
        assert "TRƯỜNG ABC" in text

    def test_mc_question_content(self):
        doc = self._render(questions=[MC_Q])
        text = _text(doc)
        assert "Câu hỏi MC?" in text

    def test_mc_options_rendered(self):
        doc = self._render(questions=[MC_Q])
        text = _text(doc)
        assert "Đáp án A" in text
        assert "Đáp án B" in text

    def test_ma_question_content(self):
        doc = self._render(questions=[MA_Q])
        text = _text(doc)
        assert "Câu hỏi MA?" in text

    def test_tf_question_answer_sheet_rendered(self):
        doc = self._render(questions=[TF_Q])
        text = _text(doc)
        assert "Đúng" in text
        assert "Sai" in text

    def test_blank_question_answer_space(self):
        doc = self._render(questions=[BLANK_Q])
        text = _text(doc)
        assert "Trả lời:" in text

    def test_sa_question_multi_line_space(self):
        doc = self._render(questions=[SA_Q])
        paragraphs_with_lines = [p for p in doc.paragraphs if "_" * 40 in p.text]
        # SA renders 3 underline lines
        assert len(paragraphs_with_lines) >= 3

    def test_no_instructions_when_disabled(self):
        cfg = ExportConfig(show_instructions=False)
        doc = self._render(config=cfg)
        text = _text(doc)
        assert "HƯỚNG DẪN LÀM BÀI" not in text

    def test_instructions_present_by_default(self):
        doc = self._render()
        text = _text(doc)
        assert "HƯỚNG DẪN LÀM BÀI" in text

    def test_no_answer_sheet_when_disabled(self):
        cfg = ExportConfig(show_answer_sheet=False)
        doc = self._render(config=cfg)
        text = _text(doc)
        assert "PHIẾU TRẢ LỜI" not in text

    def test_answer_sheet_present_by_default(self):
        doc = self._render()
        text = _text(doc)
        assert "PHIẾU TRẢ LỜI" in text

    def test_no_scoring_rules_when_disabled(self):
        cfg = ExportConfig(show_scoring_rules=False)
        doc = self._render(config=cfg)
        text = _text(doc)
        assert "QUY ĐỊNH CHẤM ĐIỂM" not in text

    def test_scoring_rules_present_by_default(self):
        doc = self._render()
        text = _text(doc)
        assert "QUY ĐỊNH CHẤM ĐIỂM" in text

    def test_question_statistics_are_hidden_by_default(self):
        doc = self._render()
        text = _text(doc)
        assert "Thống kê theo Chương" not in text
        assert "Thống kê theo CLO" not in text

    def test_question_statistics_can_be_rendered_as_supplementary_page(self):
        cfg = ExportConfig(show_question_statistics=True)
        doc = self._render(config=cfg)
        text = _text(doc)
        assert "Thống kê theo Chương" in text
        assert "Thống kê theo CLO" in text
        assert "Thống kê theo Mức độ" in text
        assert "Thống kê theo Loại câu hỏi" in text
        assert "Tỷ lệ điểm" in text
        assert "100.0%" in text
        assert "11.8%" in text

    def test_no_answer_key_when_disabled(self):
        cfg = ExportConfig(show_answer_key=False)
        doc = self._render(config=cfg)
        text = _text(doc)
        assert "ĐÁP ÁN VÀ THANG ĐIỂM" not in text

    def test_answer_key_present_by_default(self):
        doc = self._render()
        text = _text(doc)
        assert "ĐÁP ÁN VÀ THANG ĐIỂM" in text

    def test_student_info_row_uses_shorter_class_and_stt_lines(self):
        doc = self._render(questions=[MC_Q])
        text = _text(doc)
        assert "MSSV:" in text
        assert "Lớp:  __________" in text
        assert "STT:  ___" in text

    def test_question_points_are_hidden_by_default_in_exam_body(self):
        doc = self._render(questions=[MC_Q, MA_Q])
        text = _text(doc)
        assert "[1 điểm]" not in text
        assert "[2 điểm]" not in text

    def test_question_points_can_be_shown_in_exam_body_when_enabled(self):
        cfg = ExportConfig(show_question_points=True)
        doc = self._render(questions=[MC_Q, MA_Q], config=cfg)
        text = _text(doc)
        assert "[1 điểm]" in text
        assert "[2 điểm]" in text

    def test_student_signature_column_is_reduced_to_ninety_percent_width(self):
        doc = self._render(questions=[MC_Q])
        student_table = next(
            table for table in doc.tables
            if table.cell(0, 0).text.strip() == "Thông tin của sinh viên"
            and table.cell(0, 1).text.strip() == "Chữ ký của sinh viên"
        )
        left_width = _cell_width_twips(student_table.rows[0].cells[0])
        right_width = _cell_width_twips(student_table.rows[0].cells[1])
        assert right_width == pytest.approx(int(left_width * (8.55 / 12.5)), abs=2)

    def test_student_info_block_can_be_hidden(self):
        cfg = ExportConfig(
            show_instructions=False,
            show_answer_sheet=False,
            show_scoring_rules=False,
            show_answer_key=False,
        )
        cfg.print_profile = PrintProfile(show_student_info_block=False)
        doc = self._render(questions=[MC_Q], config=cfg)
        text = _text(doc)
        assert "Thông tin của sinh viên" not in text
        assert "Chữ ký của sinh viên" not in text

    def test_letter_page_size_and_custom_margins_apply(self):
        cfg = ExportConfig(
            show_instructions=False,
            show_answer_sheet=False,
            show_scoring_rules=False,
            show_answer_key=False,
        )
        cfg.print_profile = PrintProfile(
            page_size="LETTER",
            top_margin_cm=2.5,
            bottom_margin_cm=2.0,
            left_margin_cm=2.2,
            right_margin_cm=1.8,
        )
        doc = self._render(questions=[MC_Q], config=cfg)
        section = doc.sections[0]
        assert section.top_margin.cm == pytest.approx(2.5, abs=0.1)
        assert section.bottom_margin.cm == pytest.approx(2.0, abs=0.1)
        assert section.left_margin.cm == pytest.approx(2.2, abs=0.1)
        assert section.right_margin.cm == pytest.approx(1.8, abs=0.1)

    def test_cover_sheet_can_be_rendered(self):
        cfg = ExportConfig(
            show_cover_sheet=True,
            show_instructions=False,
            show_answer_sheet=False,
            show_scoring_rules=False,
            show_answer_key=False,
        )
        meta = ExamMeta(
            school="Truong ABC",
            exam_title="Giua ky",
            subject="Toan roi rac",
            course_code="MTH201",
            instructor="Nguyen Van A",
            duration_minutes=45,
        )
        doc = self._render(questions=[MC_Q], config=cfg, meta=meta)
        text = _text(doc)
        assert "GIUA KY" in text
        assert "Toan roi rac" in text
        assert "Mã học phần: MTH201" in text

    def test_watermark_text_is_written_to_header(self):
        cfg = ExportConfig(
            show_instructions=False,
            show_answer_sheet=False,
            show_scoring_rules=False,
            show_answer_key=False,
            watermark_text="NOI BO",
        )
        doc = self._render(questions=[MC_Q], config=cfg)
        header_text = "\n".join(p.text for p in doc.sections[0].header.paragraphs)
        assert "NOI BO" in header_text

    def test_cover_sheet_minimal_template_renders_internal_note(self):
        cfg = ExportConfig(
            show_cover_sheet=True,
            cover_sheet_template="minimal",
            show_instructions=False,
            show_answer_sheet=False,
            show_scoring_rules=False,
            show_answer_key=False,
        )
        meta = ExamMeta(exam_title="Giua ky")
        setattr(meta, "cover_sheet_template", "minimal")
        doc = self._render(questions=[MC_Q], config=cfg, meta=meta)
        text = _text(doc)
        assert "Tài liệu dùng cho in ấn/phát hành nội bộ." in text

    def test_question_and_option_runs_use_12pt(self):
        cfg = ExportConfig(
            show_instructions=False,
            show_answer_sheet=False,
            show_scoring_rules=False,
            show_answer_key=False,
        )
        doc = self._render(questions=[MC_Q], config=cfg)

        question_p = next(p for p in doc.paragraphs if p.text.startswith("Câu 1."))
        for run in question_p.runs:
            assert run.font.size is not None
            assert run.font.size.pt == pytest.approx(12.0)

        option_p = next(p for p in doc.paragraphs if p.text.startswith("A."))
        for run in option_p.runs:
            assert run.font.size is not None
            assert run.font.size.pt == pytest.approx(12.0)

    def test_problem_question_body_renders_latex(self):
        cfg = ExportConfig(
            show_instructions=False,
            show_answer_sheet=False,
            show_scoring_rules=False,
            show_answer_key=False,
        )
        doc = self._render(questions=[PROBLEM_Q_LATEX], config=cfg)
        text = _text(doc)
        assert "\\frac" not in text
        assert "\\sqrt" not in text
        assert "μ₀" in text
        assert "√" in text
        assert "x̄" in text or "x̄" in text

    def test_problem_question_instructions_use_problem_label(self):
        text = _text(self._render(questions=[PROBLEM_Q]))
        assert "Bài toán (Problem)" in text
        assert "Câu tự luận (Essay)" not in text

    def test_problem_question_answer_lines_match_rubric_rows(self):
        cfg = ExportConfig(
            show_instructions=False,
            show_answer_sheet=False,
            show_scoring_rules=False,
            show_answer_key=False,
        )
        doc = self._render(questions=[PROBLEM_Q], config=cfg)
        line_tables = [table for table in doc.tables if len(table.rows) == len(PROBLEM_Q["problem_rubric"])]
        assert line_tables, "Expected a ruled answer table matching the rubric row count"
        for table in line_tables:
            for row in table.rows:
                assert row.height is not None
                assert row.height.inches == pytest.approx(0.3, abs=0.02)

    def test_problem_answer_sheet_uses_single_line(self):
        cfg = ExportConfig(
            show_instructions=False,
            show_answer_sheet=True,
            show_scoring_rules=False,
            show_answer_key=False,
        )
        doc = self._render(questions=[PROBLEM_Q], config=cfg)
        text = _text(doc)
        assert "PHIẾU TRẢ LỜI" in text
        assert sum(1 for p in doc.paragraphs if p.text.startswith("Câu 1:") and "_" in p.text) == 1

    def test_common_latex_commands_render_readably(self):
        rendered = render_inline_latex_text(
            r"$\mathrm{sin}(x)+\cos(x)+\sum x+\int x+\sqrt[n]{x}+\vec{v}+\widehat{AB}+\tilde{y}+\overline{z}+\lim_{n\to\infty}$"
        )
        assert "sin(x)" in rendered
        assert "cos(x)" in rendered
        assert "∑" in rendered
        assert "∫" in rendered
        assert "n√(x)" in rendered or "√(x)" in rendered
        assert "v⃗" in rendered
        assert "AB̂" in rendered or "AB̂" in rendered
        assert "ỹ" in rendered
        assert "z̄" in rendered
        assert "lim" in rendered
        assert "→" in rendered

    def test_latex_environments_render_readably(self):
        rendered = render_inline_latex_text(
            r"$f(x)=\begin{cases}x^2 & x<0 \\ x & x\ge 0\end{cases}$"
        )
        assert r"\begin" not in rendered
        assert r"\end" not in rendered
        assert "x²" in rendered
        assert "≤" in rendered or ">=" not in rendered
        assert "if" in rendered
        matrix = render_inline_latex_text(r"$\begin{pmatrix}1 & 2 \\ 3 & 4\end{pmatrix}$")
        assert "(" in matrix and ")" in matrix
        assert "1" in matrix and "4" in matrix

    def test_additional_matrix_environments_render_readably(self):
        bmatrix = render_inline_latex_text(r"$\begin{bmatrix}1 & 2 \\ 3 & 4\end{bmatrix}$")
        vmatrix = render_inline_latex_text(r"$\begin{vmatrix}1 & 2 \\ 3 & 4\end{vmatrix}$")
        Bmatrix = render_inline_latex_text(r"$\begin{Bmatrix}1 & 2 \\ 3 & 4\end{Bmatrix}$")
        assert "[" in bmatrix and "]" in bmatrix
        assert "|" in vmatrix
        assert "⦃" in Bmatrix and "⦄" in Bmatrix


# ---------------------------------------------------------------------------
# Numbering
# ---------------------------------------------------------------------------

class TestNumbering:
    def _render(self, questions, numbering_mode):
        meta = ExamMeta(exam_title="T")
        cfg = ExportConfig(
            numbering_mode=numbering_mode,
            show_instructions=False,
            show_answer_sheet=False,
            show_scoring_rules=False,
            show_answer_key=False,
        )
        return WordRenderer().render(questions, meta, cfg)

    def test_global_numbering_continuous(self):
        """Questions should be numbered 1, 2, 3, 4 across sections."""
        doc = self._render(ALL_QUESTIONS, "global")
        text = _text(doc)
        for i in range(1, 6):
            assert f"Câu {i}." in text

    def test_per_section_numbering_restarts(self):
        """Each section should start from 1."""
        doc = self._render([MC_Q, MC_Q, MA_Q, MA_Q], "per_section")
        text = _text(doc)
        # "Câu 1." should appear at least twice (once in MC section, once in MA)
        occurrences = text.count("Câu 1.")
        assert occurrences >= 2


# ---------------------------------------------------------------------------
# Answer key correctness
# ---------------------------------------------------------------------------

class TestAnswerKey:
    def _render_key(self, questions):
        meta = ExamMeta(exam_title="T")
        cfg = ExportConfig(
            show_instructions=False,
            show_answer_sheet=False,
            show_scoring_rules=False,
            show_answer_key=True,
        )
        doc = WordRenderer().render(questions, meta, cfg)
        return _text(doc)

    def test_mc_correct_key_in_answer_key(self):
        text = self._render_key([MC_Q])
        # MC_Q has option A as correct
        assert "A" in text

    def test_ma_correct_keys_in_answer_key(self):
        text = self._render_key([MA_Q])
        # MA_Q has A and B as correct
        assert "A" in text
        assert "B" in text

    def test_blank_accepted_answers_in_key(self):
        text = self._render_key([BLANK_Q])
        assert "Hà Nội" in text

    def test_sa_accepted_answers_in_key(self):
        text = self._render_key([SA_Q])
        assert "lực hút của Trái Đất" in text

    def test_problem_rubric_in_key(self):
        text = self._render_key([PROBLEM_Q])
        assert "Nội dung đáp án" in text
        assert "B1" in text
        assert "Tính thống kê kiểm định" in text
        assert "TỔNG" in text
        assert "6" in text

    def test_problem_template_name_is_shown_in_answer_key(self):
        text = self._render_key([PROBLEM_Q_TEMPLATE])
        assert "Mẫu rubric: Mẫu kiểm định chuẩn" in text

    def test_can_render_standalone_answer_key_document(self):
        meta = ExamMeta(exam_title="T")
        cfg = ExportConfig(group_by_type=True, numbering_mode="global")
        doc = WordRenderer().render_answer_key_document([MC_Q, BLANK_Q], meta, cfg)
        text = _text(doc)
        assert "ĐÁP ÁN VÀ THANG ĐIỂM" in text
        assert "Câu 1." in text
        assert "Hà Nội" in text

    def test_typed_problem_snapshot_preserves_rubric(self):
        meta = ExamMeta(exam_title="T")
        cfg = ExportConfig(
            show_instructions=False,
            show_answer_sheet=False,
            show_scoring_rules=False,
            show_answer_key=True,
        )
        doc = WordRenderer().render([ExportQuestionSnapshot.from_dict(PROBLEM_Q)], meta, cfg)
        text = _text(doc)
        assert "B1" in text
        assert "Nội dung đáp án" in text

    def test_problem_rubric_long_content_is_fully_rendered(self):
        text = self._render_key([PROBLEM_Q_LONG])
        assert "Phân tích đề bài rất dài để đảm bảo nội dung đầy đủ" in text
        assert "Xuống dòng kiểm tra" in text
        assert "Vẫn phải hiện đủ nội dung ở ô đáp án mẫu." in text or "Và vẫn phải hiện đủ nội dung ở ô đáp án mẫu." in text
        assert "Kết luận cuối cùng" in text

    def test_problem_rubric_latex_is_rendered_in_key(self):
        text = self._render_key([PROBLEM_Q_LATEX])
        assert "\\frac" not in text
        assert "\\sqrt" not in text
        assert "μ₀" in text
        assert "√" in text
        assert "x̄" in text or "x̄" in text

    def test_problem_rubric_raw_latex_mode_keeps_commands(self):
        meta = ExamMeta(exam_title="T")
        cfg = ExportConfig(
            show_instructions=False,
            show_answer_sheet=False,
            show_scoring_rules=False,
            show_answer_key=True,
            raw_latex_answer_key=True,
        )
        doc = WordRenderer().render([PROBLEM_Q_LATEX], meta, cfg)
        text = _text(doc)
        assert "\\frac{\\bar{x}-\\mu_0}{s/\\sqrt{n}}" in text
        assert "\\alpha" not in text


# ---------------------------------------------------------------------------
# Total score
# ---------------------------------------------------------------------------

class TestTotalScore:
    def test_total_score_in_scoring_rules(self):
        # Total = 1+2+1+1.5+3 = 8.5
        meta = ExamMeta(exam_title="T")
        cfg = ExportConfig(
            show_instructions=False,
            show_answer_sheet=False,
            show_scoring_rules=True,
            show_answer_key=False,
        )
        doc = WordRenderer().render(ALL_QUESTIONS, meta, cfg)
        text = _text(doc)
        assert "8.5" in text

    def test_total_score_in_answer_key(self):
        meta = ExamMeta(exam_title="T")
        cfg = ExportConfig(
            show_instructions=False,
            show_answer_sheet=False,
            show_scoring_rules=False,
            show_answer_key=True,
        )
        doc = WordRenderer().render(ALL_QUESTIONS, meta, cfg)
        text = _text(doc)
        assert "8.5" in text


# ---------------------------------------------------------------------------
# Instructions per type
# ---------------------------------------------------------------------------

class TestInstructions:
    def test_mc_instructions_present(self):
        meta = ExamMeta(exam_title="T")
        cfg = ExportConfig(show_instructions=True, show_answer_sheet=False,
                           show_scoring_rules=False, show_answer_key=False)
        doc = WordRenderer().render([MC_Q], meta, cfg)
        text = _text(doc)
        assert "Multiple Choice" in text

    def test_sa_instructions_not_shown_when_type_absent(self):
        meta = ExamMeta(exam_title="T")
        cfg = ExportConfig(show_instructions=True, show_answer_sheet=False,
                           show_scoring_rules=False, show_answer_key=False)
        doc = WordRenderer().render([MC_Q], meta, cfg)
        text = _text(doc)
        assert "Trả lời ngắn" not in text


# ---------------------------------------------------------------------------
# build_output_path
# ---------------------------------------------------------------------------

class TestBuildOutputPath:
    def test_returns_docx_extension(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = build_output_path("My Exam", Path(tmpdir))
        assert path.suffix == ".docx"

    def test_title_in_filename(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = build_output_path("My Exam", Path(tmpdir))
        assert "My_Exam" in path.name or "My Exam" in path.name

    def test_timestamp_in_filename(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = build_output_path("Exam", Path(tmpdir))
        # Timestamp pattern like 2026-04-21_143022
        assert re.search(r"\d{4}-\d{2}-\d{2}_\d{6}", path.name)

    def test_sanitises_special_characters(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = build_output_path('Exam <> "test"', Path(tmpdir))
        assert "<" not in path.name
        assert ">" not in path.name
        assert '"' not in path.name

    def test_creates_exports_dir(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            exports = Path(tmpdir) / "nested" / "exports"
            build_output_path("Exam", exports)
            assert exports.exists()

    def test_empty_title_fallback(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = build_output_path("", Path(tmpdir))
        assert "exam" in path.name.lower()

"""Word document renderer for exam export.

Accepts a list of typed export snapshots (or legacy dict snapshots)
plus exam metadata, and generates
a clean .docx exam document ready for printing.

No database access occurs here. All data must be supplied by the caller.

Supported question types: MC, MA, TF, BLANK, SA, ES, PR

Generated sections (each controlled by ExportConfig):
    1. Exam header (school, department, subject, title, etc.)
    2. Instructions
    3. Question body (grouped per section A/B/C/D or flat)
    4. Answer sheet
    5. Scoring rules
    6. Answer key and score table
"""
from __future__ import annotations

import copy
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from core.utils.blank_rendering import render_blank_placeholders
from core.utils.latex_rendering import render_inline_latex_text

# ---------------------------------------------------------------------------
# Public data classes
# ---------------------------------------------------------------------------

@dataclass
class ExamMeta:
    """Metadata for the exam header."""
    school: str = ""
    department: str = ""
    instructor: str = ""
    subject: str = ""
    course_code: str = ""
    exam_title: str = "BÀI KIỂM TRA"
    exam_type: str = ""
    duration_minutes: int = 0
    note: str = ""


@dataclass
class ExportConfig:
    """Controls which optional sections to include in the export."""
    show_instructions: bool = True
    show_answer_sheet: bool = True
    show_answer_key: bool = True
    show_scoring_rules: bool = True
    show_question_points: bool = False
    show_question_statistics: bool = False
    # "global" = continuous numbering across all questions
    # "per_section" = restart numbering in each section
    numbering_mode: str = "global"
    # "grouped" = separate Part A/B/C/D by question type
    # "flat"    = all questions in one list
    group_by_type: bool = True
    show_cover_sheet: bool = False
    split_answer_key_file: bool = False
    raw_latex_answer_key: bool = False
    watermark_text: str = ""
    watermark_preset: str = "custom"
    cover_sheet_template: str = "standard"
    answer_key_naming_policy: str = "suffix"
    # CRQ questions appended as last section
    # each dict: {"number": int, "score": float}
    essay_questions: list[dict] = field(default_factory=list)


@dataclass
class PrintProfile:
    """Controls page layout and printable header blocks."""

    page_size: str = "A4"
    top_margin_cm: float = 1.5
    bottom_margin_cm: float = 1.5
    left_margin_cm: float = 2.0
    right_margin_cm: float = 1.5
    show_student_info_block: bool = True


@dataclass
class ExportQuestionSnapshot:
    """Typed snapshot payload used by WordRenderer export flow."""

    type: str
    content: str
    point_value: float = 1.0
    hint: str = ""
    explanation: str = ""
    options: list = field(default_factory=list)
    accepted_answers: list | None = None
    question_variant: str = ""
    question_family: str = ""
    crq_subtype: str = ""
    crq_rubric: list | None = None
    crq_template_id: int | None = None
    crq_template_name: str = ""
    problem_rubric: list | None = None
    problem_template_id: int | None = None
    problem_template_name: str = ""
    difficulty: str = ""
    learning_outcome_code: str = ""
    category: str = ""

    @classmethod
    def from_dict(cls, data: dict) -> ExportQuestionSnapshot:
        return cls(
            type=data.get("type", "MC"),
            content=data.get("content", ""),
            point_value=data.get("point_value", 1.0),
            hint=data.get("hint") or "",
            explanation=data.get("explanation") or "",
            options=data.get("options") or [],
            accepted_answers=data.get("accepted_answers"),
            question_variant=data.get("question_variant") or "",
            question_family=data.get("question_family") or "",
            crq_subtype=data.get("crq_subtype") or "",
            crq_rubric=data.get("crq_rubric"),
            crq_template_id=data.get("crq_template_id"),
            crq_template_name=data.get("crq_template_name") or "",
            problem_rubric=data.get("problem_rubric"),
            problem_template_id=data.get("problem_template_id"),
            problem_template_name=data.get("problem_template_name") or "",
            difficulty=data.get("difficulty") or "",
            learning_outcome_code=data.get("learning_outcome_code") or "",
            category=data.get("category") or "",
        )

    def to_dict(self) -> dict:
        return {
            "type": self.type,
            "content": self.content,
            "point_value": self.point_value,
            "hint": self.hint,
            "explanation": self.explanation,
            "options": self.options,
            "accepted_answers": self.accepted_answers,
            "question_variant": self.question_variant,
            "question_family": self.question_family,
            "crq_subtype": self.crq_subtype,
            "crq_rubric": self.crq_rubric,
            "crq_template_id": self.crq_template_id,
            "crq_template_name": self.crq_template_name,
            "problem_rubric": self.problem_rubric,
            "problem_template_id": self.problem_template_id,
            "problem_template_name": self.problem_template_name,
            "difficulty": self.difficulty,
            "learning_outcome_code": self.learning_outcome_code,
            "category": self.category,
        }


# Internal type → display label mapping
_TYPE_LABELS: dict[str, str] = {
    "MC": "Multiple Choice",
    "MA": "Multiple Answer",
    "TF": "Đúng/Sai",
    "BLANK": "Điền vào chỗ trống",
    "SA": "Trả lời ngắn",
    "ES": "CRQ - Tự luận",
    "PR": "CRQ - Bài toán",
}

_TYPE_STATS_LABELS: dict[str, str] = {
    "MC": "Trắc nghiệm 1 đáp án",
    "MA": "Trắc nghiệm nhiều đáp án",
    "TF": "Đúng/Sai",
    "BLANK": "Điền vào chỗ trống",
    "SA": "Trả lời ngắn",
    "ES": "CRQ - Tự luận",
    "PR": "CRQ - Bài toán",
}

_DIFFICULTY_LABELS: dict[str, str] = {
    "easy": "Nhớ",
    "medium": "Hiểu",
    "hard": "Vận dụng",
    "Nhớ": "Nhớ",
    "Hiểu": "Hiểu",
    "Vận dụng": "Vận dụng",
    "Phân tích": "Phân tích",
    "Đánh giá": "Đánh giá",
    "Sáng tạo": "Sáng tạo",
}

_SECTION_LETTERS = "ABCDEFGH"

# Instructions text per question type
_INSTRUCTIONS: dict[str, str] = {
    "MC": (
        "Câu hỏi trắc nghiệm một đáp án (Multiple Choice): "
        "Chọn MỘT đáp án đúng nhất bằng cách khoanh tròn vào chữ cái tương ứng."
    ),
    "MA": (
        "Câu hỏi trắc nghiệm nhiều đáp án (Multiple Answer): "
        "Chọn TẤT CẢ các đáp án đúng bằng cách khoanh tròn hoặc đánh dấu vào chữ cái tương ứng."
    ),
    "TF": (
        "Câu hỏi Đúng/Sai (True/False): "
        "Chọn một đáp án đúng trong hai lựa chọn Đúng hoặc Sai."
    ),
    "BLANK": (
        "Câu điền vào chỗ trống (Blank): "
        "Điền từ hoặc cụm từ thích hợp vào chỗ trống trong câu."
    ),
    "SA": (
        "Câu trả lời ngắn (Short Answer): "
        "Trả lời ngắn gọn, rõ ràng vào phần dành sẵn bên dưới mỗi câu."
    ),
    "ES": (
        "Câu CRQ - Tự luận (Essay): "
        "Trình bày câu trả lời theo hướng dẫn hoặc dàn ý được yêu cầu."
    ),
    "PR": (
        "Câu CRQ - Bài toán (Problem): "
        "Trình bày lời giải theo hướng dẫn hoặc rubric được yêu cầu."
    ),
}

# ---------------------------------------------------------------------------
# Main renderer class
# ---------------------------------------------------------------------------

class WordRenderer:
    """Renders a list of question snapshots into a .docx exam document.

    Usage::

        renderer = WordRenderer()
        doc = renderer.render(questions, meta, config)
        doc.save("exam.docx")

    Parameters
    ----------
    questions:
        List of snapshot dicts as produced by
        ``QuestionSelector.build_snapshots()``.  Each dict must have at
        minimum the keys: ``type``, ``content``, ``point_value``.
        MC/MA also need ``options`` (list of ``{key, text, is_correct}``).
        BLANK/SA also need ``accepted_answers`` (list[str]).
    meta:
        ``ExamMeta`` instance with header information.
    config:
        ``ExportConfig`` instance controlling optional sections.
    """

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def render(
        self,
        questions: list[ExportQuestionSnapshot | dict],
        meta: ExamMeta,
        config: ExportConfig,
    ) -> Document:
        """Build and return a ``docx.Document`` object.

        The caller is responsible for saving the document to disk.
        """
        normalized_questions = self._normalize_questions(questions)

        doc = Document()
        self._set_default_font(doc, "Times New Roman")
        self._apply_print_profile(doc, self._build_print_profile(config))

        print_profile = self._build_print_profile(config)
        self._apply_watermark(doc, config.watermark_text)

        if config.show_cover_sheet:
            self._render_cover_sheet(doc, meta)
            doc.add_page_break()

        self._render_header(doc, meta, print_profile)

        if config.show_instructions:
            self._render_instructions(doc, normalized_questions)

        grouped = self._group_questions(normalized_questions, config)
        self._render_questions(doc, grouped, config)

        if config.essay_questions:
            essay_letter_idx = len(grouped) if config.group_by_type else 0
            essay_letter = (
                _SECTION_LETTERS[essay_letter_idx]
                if essay_letter_idx < len(_SECTION_LETTERS)
                else ""
            )
            self._render_essay_section(doc, config.essay_questions, essay_letter, config)

        if (
            config.show_answer_sheet
            or config.show_scoring_rules
            or config.show_answer_key
            or config.show_question_statistics
        ):
            # Close the question section — SECTIONPAGES in the header will reflect
            # only question pages, not the supplementary pages that follow.
            self._add_next_page_section_break(doc)
            # Supplementary section must not show any page-number header.
            self._clear_supplementary_header(doc)
            first_supp = True

            if config.show_answer_sheet:
                if not first_supp:
                    doc.add_page_break()
                first_supp = False
                self._render_answer_sheet(doc, grouped, config)

            if config.show_scoring_rules:
                if not first_supp:
                    doc.add_page_break()
                first_supp = False
                self._render_scoring_rules(doc, normalized_questions, config)

            if config.show_question_statistics:
                if not first_supp:
                    doc.add_page_break()
                first_supp = False
                self._render_question_statistics(doc, normalized_questions, config)

            if config.show_answer_key:
                if not first_supp:
                    doc.add_page_break()
                self._render_answer_key(doc, grouped, config)

        return doc

    def render_answer_key_document(
        self,
        questions: list[ExportQuestionSnapshot | dict],
        meta: ExamMeta,
        config: ExportConfig,
    ) -> Document:
        """Render a standalone answer-key document for teacher use."""
        normalized_questions = self._normalize_questions(questions)
        grouped = self._group_questions(normalized_questions, config)

        doc = Document()
        self._set_default_font(doc, "Times New Roman")
        self._apply_print_profile(doc, self._build_print_profile(config))
        self._apply_watermark(doc, config.watermark_text)

        key_meta = ExamMeta(**vars(meta))
        key_meta.exam_title = f"{meta.exam_title} - ĐÁP ÁN"
        self._render_header(doc, key_meta, self._build_print_profile(config))
        self._render_answer_key(doc, grouped, config)
        return doc

    @staticmethod
    def _apply_paragraph_spacing(paragraph, *, before_pt: float = 3.0, after_pt: float = 3.0) -> None:
        fmt = paragraph.paragraph_format
        fmt.space_before = Pt(before_pt)
        fmt.space_after = Pt(after_pt)

    @staticmethod
    def _render_key_text(text: str, *, raw_latex: bool) -> str:
        value = str(text or "")
        return value if raw_latex else render_inline_latex_text(value)

    @staticmethod
    def _normalize_questions(
        questions: list[ExportQuestionSnapshot | dict],
    ) -> list[dict]:
        """Convert typed snapshots to legacy dict shape for renderer internals."""
        normalized: list[dict] = []
        for q in questions:
            if isinstance(q, ExportQuestionSnapshot):
                normalized.append(q.to_dict())
            else:
                normalized.append(q)
        return normalized

    # ------------------------------------------------------------------
    # Page setup
    # ------------------------------------------------------------------

    def _build_print_profile(self, config: ExportConfig) -> PrintProfile:
        profile = getattr(config, "print_profile", None)
        if isinstance(profile, PrintProfile):
            return profile
        return PrintProfile(
            page_size=getattr(config, "page_size", "A4"),
            top_margin_cm=float(getattr(config, "top_margin_cm", 1.5)),
            bottom_margin_cm=float(getattr(config, "bottom_margin_cm", 1.5)),
            left_margin_cm=float(getattr(config, "left_margin_cm", 2.0)),
            right_margin_cm=float(getattr(config, "right_margin_cm", 1.5)),
            show_student_info_block=bool(getattr(config, "show_student_info_block", True)),
        )

    def _apply_print_profile(self, doc: Document, profile: PrintProfile) -> None:
        section = doc.sections[0]
        self._set_page_size(section, profile.page_size)
        section.top_margin = Cm(profile.top_margin_cm)
        section.bottom_margin = Cm(profile.bottom_margin_cm)
        section.left_margin = Cm(profile.left_margin_cm)
        section.right_margin = Cm(profile.right_margin_cm)

    def _set_page_size(self, section, page_size: str) -> None:
        size = page_size.upper().strip()
        if size == "LETTER":
            section.page_width = Inches(8.5)
            section.page_height = Inches(11.0)
            return
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    def _apply_watermark(self, doc: Document, watermark_text: str) -> None:
        text = self._resolve_watermark_text(watermark_text)
        if not text:
            return
        header = doc.sections[0].header
        p = header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(26)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0xC8, 0xC8, 0xC8)

    def _resolve_watermark_text(self, watermark_text: str) -> str:
        return watermark_text.strip()

    def _set_default_font(self, doc: Document, font_name: str) -> None:
        """Apply *font_name* as the document-wide default for all runs.

        Sets both the 'Normal' paragraph style (which most styles inherit from)
        and the low-level ``w:docDefaults`` XML element so that runs without an
        explicit font also pick up the correct typeface.
        """
        # 1. Normal style – all paragraph styles inherit from this
        doc.styles["Normal"].font.name = font_name
        doc.styles["Normal"].font.size = Pt(12)

        # 2. Document-level rPrDefault – catches runs that bypass style inheritance
        styles_el = doc.styles.element
        docDefaults = styles_el.find(qn("w:docDefaults"))
        if docDefaults is None:
            return
        rPrDefault = docDefaults.find(qn("w:rPrDefault"))
        if rPrDefault is None:
            return
        rPr = rPrDefault.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            rPrDefault.append(rPr)
        # Remove any existing rFonts element before inserting
        for existing in rPr.findall(qn("w:rFonts")):
            rPr.remove(existing)
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
        rFonts.set(qn("w:cs"), font_name)
        rPr.insert(0, rFonts)
        for existing in rPr.findall(qn("w:sz")):
            rPr.remove(existing)
        for existing in rPr.findall(qn("w:szCs")):
            rPr.remove(existing)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "24")
        sz_cs = OxmlElement("w:szCs")
        sz_cs.set(qn("w:val"), "24")
        rPr.append(sz)
        rPr.append(sz_cs)
    # ------------------------------------------------------------------
    # Header
    # ------------------------------------------------------------------

    def _render_header(self, doc: Document, meta: ExamMeta, print_profile: PrintProfile) -> None:
        """Header matching the Vietnamese standard exam template.

        Layout::

            [Word page header: Trang X / Y, top right]

            [Borderless 2-col info table]
            TRƯỜNG...          | BÀI KIỂM TRA GIỮA KỲ
            Khoa/Đơn vị...     | Môn học (Mã HP)
            Cán bộ giảng dạy.. | Hình thức: ... – Thời lượng: __ phút

            [Bordered 2-col student table]
            | Thông tin của sinh viên   | Chữ ký của sinh viên |
            | Họ & tên: ____________    |                      |
            | MSSV: __ Lớp: __ STT: __ |                      |

            Lưu ý: <meta.note>
        """
        # ── Page number in Word built-in header (top right) ──────────────
        self._set_page_number_header(doc)

        # ── Info block (borderless 2-column table) ───────────────────────
        INFO_LEFT_CM = 9.5
        INFO_RIGHT_CM = 9.5
        info_tbl = doc.add_table(rows=1, cols=2)
        self._clear_table_borders(info_tbl)
        L = info_tbl.cell(0, 0)
        R = info_tbl.cell(0, 1)
        self._set_cell_width(L, INFO_LEFT_CM)
        self._set_cell_width(R, INFO_RIGHT_CM)

        # Left cell: school → department → instructor
        lp = L.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if meta.school:
            r = lp.add_run(meta.school.upper())
            r.bold = True
            r.font.size = Pt(13)
        if meta.department:
            p = L.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(meta.department)
            r.bold = True
            r.font.size = Pt(13)
        if meta.instructor:
            p = L.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f"Cán bộ giảng dạy: {meta.instructor}")
            r.italic = True
            r.font.size = Pt(12)

        # Right cell: title → subject/code → format/duration
        rp = R.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if meta.exam_title:
            r = rp.add_run(meta.exam_title.upper())
            r.bold = True
            r.font.size = Pt(13)
        if meta.subject or meta.course_code:
            parts = [meta.subject] if meta.subject else []
            if meta.course_code:
                parts.append(f"({meta.course_code})")
            p = R.add_paragraph(" ".join(parts))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(12)
        fmt_line = self._build_format_duration_line(meta)
        if fmt_line:
            p = R.add_paragraph(fmt_line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(12)

        doc.add_paragraph("")  # gap

        if print_profile.show_student_info_block:
            self._render_student_info_block(doc)

        # ── Note paragraph ───────────────────────────────────────────────
        doc.add_paragraph("")
        if meta.note:
            p = doc.add_paragraph()
            r1 = p.add_run("Lưu ý: ")
            r1.bold = True
            r1.font.size = Pt(12)
            r2 = p.add_run(meta.note)
            r2.font.size = Pt(12)
        doc.add_paragraph("")  # spacer before section body

    def _render_cover_sheet(self, doc: Document, meta: ExamMeta) -> None:
        """Render a simple teacher-facing cover sheet before the exam body."""
        template = getattr(meta, "cover_sheet_template", "standard")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run((meta.school or "").upper())
        run.bold = True
        run.font.size = Pt(16)

        if meta.department and template == "standard":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(meta.department)
            run.bold = True
            run.font.size = Pt(14)

        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(meta.exam_title.upper())
        run.bold = True
        run.font.size = Pt(20)

        if meta.subject:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(meta.subject)
            run.font.size = Pt(16)

        if meta.course_code:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Mã học phần: {meta.course_code}")
            run.font.size = Pt(13)

        if meta.instructor and template == "standard":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Cán bộ giảng dạy: {meta.instructor}")
            run.font.size = Pt(13)

        duration = meta.duration_minutes if meta.duration_minutes else "__"
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Thời lượng: {duration} phút")
        run.font.size = Pt(13)

        if meta.note:
            doc.add_paragraph("")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(meta.note)
            run.italic = True
            run.font.size = Pt(12)

        if template == "minimal":
            doc.add_paragraph("")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("Tài liệu dùng cho in ấn/phát hành nội bộ.")
            run.italic = True
            run.font.size = Pt(11)

    def _render_student_info_block(self, doc: Document) -> None:
        """Render the bordered student information block below exam metadata."""
        STU_LEFT_CM = 12.5
        STU_RIGHT_CM = 8.55
        stu_tbl = doc.add_table(rows=3, cols=2)
        stu_tbl.style = "Table Grid"
        for row in stu_tbl.rows:
            self._set_cell_width(row.cells[0], STU_LEFT_CM)
            self._set_cell_width(row.cells[1], STU_RIGHT_CM)

        p = stu_tbl.cell(0, 0).paragraphs[0]
        r = p.add_run("Thông tin của sinh viên")
        r.bold = True
        r.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = stu_tbl.cell(0, 1).paragraphs[0]
        r = p.add_run("Chữ ký của sinh viên")
        r.bold = True
        r.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = stu_tbl.cell(1, 0).paragraphs[0]
        p.add_run("Họ & tên:  " + "_" * 44)
        p.runs[0].font.size = Pt(12)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

        p = stu_tbl.cell(2, 0).paragraphs[0]
        p.add_run("MSSV:  " + "_" * 16 + "   Lớp:  " + "_" * 10 + "   STT:  " + "_" * 3)
        p.runs[0].font.size = Pt(12)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

        stu_tbl.cell(1, 1).merge(stu_tbl.cell(2, 1))
        self._hide_cell_border(stu_tbl.cell(1, 0), "bottom")
        self._hide_cell_border(stu_tbl.cell(2, 0), "top")

    # ------------------------------------------------------------------
    # Header helpers
    # ------------------------------------------------------------------

    def _set_page_number_header(self, doc: Document) -> None:
        """Insert 'Trang X / Y' page-number field in the Word document header (top right)."""
        header = doc.sections[0].header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run = p.add_run("Trang ")
        run.font.size = Pt(12)
        self._add_field_code(p, "PAGE")
        run2 = p.add_run(" / ")
        run2.font.size = Pt(12)
        self._add_field_code(p, "SECTIONPAGES")

    def _clear_supplementary_header(self, doc: Document) -> None:
        """Give the supplementary section a blank, independent header (no page numbers).

        After ``_add_next_page_section_break`` deep-copies ``sectPr``, both the
        question section (inline sectPr) and the supplementary section (body
        sectPr) share the **same** header relationship ID.  We must give the
        supplementary section its own new empty header part so Word does not
        inherit the question-section page-number header on those pages.

        ``Section.header_is_linked_to_previous`` does NOT exist in
        python-docx 1.2.0 — using that attribute would be a silent no-op.
        Instead we:
        1. Create a fresh empty header part via ``doc.part.add_header_part()``.
        2. Replace the ``<w:headerReference r:id>`` in the body sectPr with the
           new rId, leaving the inline sectPr's original rId untouched.
        """
        R_ID = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
        # Create a new empty header part and register it in doc.part relationships
        _empty_hdr_part, new_rId = doc.part.add_header_part()

        # The supplementary section is always the body's final sectPr (sections[-1])
        supp_sectPr = doc.sections[-1]._sectPr

        # Replace every existing headerReference in the supplementary sectPr
        for href in supp_sectPr.findall(qn("w:headerReference")):
            supp_sectPr.remove(href)

        href_el = OxmlElement("w:headerReference")
        href_el.set(qn("w:type"), "default")
        href_el.set(R_ID, new_rId)
        supp_sectPr.append(href_el)

    def _add_field_code(self, paragraph, field_name: str) -> None:
        """Append a Word field code run (PAGE, NUMPAGES, etc.) to *paragraph*."""
        run = paragraph.add_run()
        run.font.size = Pt(12)
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_begin)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f" {field_name} "
        run._r.append(instr)
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_end)

    def _hide_cell_border(self, cell, *sides: str) -> None:
        """Set specific border *sides* of *cell* to 'nil' (invisible).

        Accepted side names: ``top``, ``bottom``, ``left``, ``right``,
        ``insideH``, ``insideV``.
        """
        tc = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            tc.insert(0, tcPr)
        tcBorders = tcPr.find(qn("w:tcBorders"))
        if tcBorders is None:
            tcBorders = OxmlElement("w:tcBorders")
            tcPr.append(tcBorders)
        for side in sides:
            for old in tcBorders.findall(qn(f"w:{side}")):
                tcBorders.remove(old)
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "nil")
            tcBorders.append(el)

    def _clear_table_borders(self, table) -> None:
        """Remove all visible borders from a table (make it invisible)."""
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        existing = tblPr.find(qn("w:tblBorders"))
        if existing is not None:
            tblPr.remove(existing)
        tblBorders = OxmlElement("w:tblBorders")
        for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{name}")
            el.set(qn("w:val"), "none")
            el.set(qn("w:sz"), "0")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "auto")
            tblBorders.append(el)
        tblPr.append(tblBorders)

    def _set_cell_width(self, cell, width_cm: float) -> None:
        """Set the width of a table cell (in centimetres) via raw XML."""
        tc = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            tc.insert(0, tcPr)
        existing = tcPr.find(qn("w:tcW"))
        if existing is not None:
            tcPr.remove(existing)
        tcW = OxmlElement("w:tcW")
        # 1 cm = 360 000 EMU; 1 twip = 635 EMU → twips = Cm(x) / 635
        twips = int(Cm(width_cm) / 635)
        tcW.set(qn("w:w"), str(twips))
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)

    def _build_format_duration_line(self, meta: ExamMeta) -> str:
        """Return the 'Hình thức: X – Thời lượng: Y phút' header line."""
        parts = []
        if meta.exam_type:
            parts.append(f"Hình thức: {meta.exam_type}")
        dur = str(meta.duration_minutes) if meta.duration_minutes else "__"
        parts.append(f"Thời lượng: {dur} phút")
        return " \u2013 ".join(parts)  # en-dash separator

    def _add_next_page_section_break(self, doc: Document) -> None:
        """Insert a 'next page' Word section break after the question body.

        Using SECTIONPAGES in the document header means the counter resets
        to the page count of this first section only (question pages),
        excluding the answer sheet, scoring rules, and answer key pages.
        """
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        pPr = p._p.get_or_add_pPr()
        # Deep-copy the main sectPr to preserve page size and margins
        main_sectPr = doc.sections[0]._sectPr
        sectPr = copy.deepcopy(main_sectPr)
        # Ensure the type element is set to nextPage
        type_el = sectPr.find(qn("w:type"))
        if type_el is None:
            type_el = OxmlElement("w:type")
            sectPr.insert(0, type_el)
        type_el.set(qn("w:val"), "nextPage")
        pPr.append(sectPr)

    # ------------------------------------------------------------------
    # Instructions
    # ------------------------------------------------------------------

    def _render_instructions(self, doc: Document, questions: list[dict]) -> None:
        types_present = self._types_present(questions)
        if not types_present:
            return

        p = doc.add_paragraph("HƯỚNG DẪN LÀM BÀI")
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(12)
        p.runs[0].underline = True
        self._apply_paragraph_spacing(p)

        for qtype in ["MC", "MA", "TF", "BLANK", "SA", "ES", "PR"]:
            if qtype in types_present and qtype in _INSTRUCTIONS:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(_INSTRUCTIONS[qtype])
                p.runs[0].font.size = Pt(12)
                self._apply_paragraph_spacing(p)

        doc.add_paragraph("")

    # ------------------------------------------------------------------
    # Question body
    # ------------------------------------------------------------------

    def _group_questions(
        self,
        questions: list[dict],
        config: ExportConfig,
    ) -> list[tuple[str, str, list[dict]]]:
        """Return list of (section_letter, section_title, questions).

        If group_by_type is False, returns a single section with all questions.
        """
        if not config.group_by_type:
            return [("", "", list(questions))]

        # Group by type in canonical order
        order = ["MC", "MA", "TF", "BLANK", "SA", "ES", "PR"]
        buckets: dict[str, list[dict]] = {t: [] for t in order}
        for q in questions:
            qtype = q.get("type", "MC")
            if qtype in buckets:
                buckets[qtype].append(q)

        result = []
        letter_idx = 0
        for qtype in order:
            if buckets[qtype]:
                letter = _SECTION_LETTERS[letter_idx]
                title_label = self._section_title_label(qtype, buckets[qtype])
                title = f"Phần {letter}. {title_label}"
                result.append((letter, title, buckets[qtype]))
                letter_idx += 1
        return result

    def _render_questions(
        self,
        doc: Document,
        grouped: list[tuple[str, str, list[dict]]],
        config: ExportConfig,
    ) -> None:
        global_counter = 0

        for _letter, title, qs in grouped:
            if title:
                p = doc.add_paragraph(title)
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(12)
                run.underline = True
                self._apply_paragraph_spacing(p)

            for section_counter, q in enumerate(qs, start=1):
                global_counter += 1
                num = (
                    global_counter
                    if config.numbering_mode == "global"
                    else section_counter
                )
                self._render_single_question(doc, q, num, config)

            doc.add_paragraph("")

    def _render_single_question(
        self,
        doc: Document,
        q: dict,
        number: int,
        config: ExportConfig,
    ) -> None:
        qtype = q.get("type", "MC")
        content = render_inline_latex_text(
            render_blank_placeholders(str(q.get("content", "")))
        )
        points = q.get("point_value", 1.0)

        # Question stem
        p = doc.add_paragraph()
        run = p.add_run(f"Câu {number}.")
        run.bold = True
        run.font.size = Pt(12)
        self._apply_paragraph_spacing(p)
        p.add_run(f" {content}")
        p.runs[-1].font.size = Pt(12)
        if config.show_question_points:
            p.add_run(f"  [{points:g} điểm]")
            p.runs[-1].font.size = Pt(12)
            p.runs[-1].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        if qtype in ("MC", "MA", "TF"):
            self._render_mc_ma_options(doc, q)
        elif qtype == "BLANK":
            self._render_blank_answer_space(doc)
        elif qtype == "SA":
            self._render_sa_answer_space(doc)
        elif qtype in ("ES", "PR"):
            self._render_problem_answer_space(doc, q)

    def _render_mc_ma_options(self, doc: Document, q: dict) -> None:
        options = q.get("options", [])
        for opt in options:
            p = doc.add_paragraph()
            run = p.add_run(f"{opt.get('key', '')}. ")
            run.bold = True
            run.font.size = Pt(12)
            p.add_run(render_inline_latex_text(str(opt.get("text", ""))))
            p.runs[-1].font.size = Pt(12)
            p.paragraph_format.left_indent = Cm(1.0)
            self._apply_paragraph_spacing(p)

    def _render_blank_answer_space(self, doc: Document) -> None:
        p = doc.add_paragraph("Trả lời: " + "_" * 40)
        p.runs[0].font.size = Pt(12)
        p.paragraph_format.left_indent = Cm(1.0)
        self._apply_paragraph_spacing(p)

    def _render_sa_answer_space(self, doc: Document) -> None:
        p = doc.add_paragraph("Trả lời:")
        p.runs[0].font.size = Pt(12)
        self._apply_paragraph_spacing(p)
        for _ in range(3):
            p2 = doc.add_paragraph("_" * 80)
            p2.runs[0].font.size = Pt(12)
            p2.paragraph_format.left_indent = Cm(1.0)
            self._apply_paragraph_spacing(p2)

    def _render_problem_answer_space(self, doc: Document, question: dict) -> None:
        line_count = self._problem_answer_line_count(question)
        p = doc.add_paragraph("Trả lời:")
        p.runs[0].font.size = Pt(12)
        self._apply_paragraph_spacing(p)
        table = doc.add_table(rows=line_count, cols=1)
        table.style = "Table Grid"
        table.autofit = False
        self._set_cell_width(table.cell(0, 0), 17.2)
        for row in table.rows:
            row.height = Inches(0.3)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            cell = row.cells[0]
            self._hide_cell_border(cell, "top", "left", "right")
            cell.paragraphs[0].paragraph_format.space_before = Pt(3)
            cell.paragraphs[0].paragraph_format.space_after = Pt(3)
            cell.paragraphs[0].paragraph_format.left_indent = Cm(0.85)
            cell.paragraphs[0].add_run("")
        doc.add_paragraph("")

    def _render_essay_section(
        self,
        doc: Document,
        essay_questions: list[dict],
        section_letter: str,
        config: ExportConfig,
    ) -> None:
        """Render CRQ section as the last part of the question body."""
        title = f"Phần {section_letter}. CRQ" if section_letter else "CRQ"
        p = doc.add_paragraph(title)
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(12)
        run.underline = True
        self._apply_paragraph_spacing(p)

        for num, eq in enumerate(essay_questions, start=1):
            score = eq.get("score", 1.0)
            p = doc.add_paragraph()
            run = p.add_run(f"Câu {num}.")
            run.bold = True
            run.font.size = Pt(12)
            self._apply_paragraph_spacing(p)
            if config.show_question_points:
                p.add_run(f"  [{score:g} điểm]")
                p.runs[-1].font.size = Pt(12)
                p.runs[-1].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            self._apply_paragraph_spacing(p)
            for _ in range(6):
                p2 = doc.add_paragraph("_" * 80)
                p2.runs[0].font.size = Pt(12)
                p2.paragraph_format.left_indent = Cm(1.0)
                self._apply_paragraph_spacing(p2)
            doc.add_paragraph("")

    # ------------------------------------------------------------------
    # Answer sheet
    # ------------------------------------------------------------------

    def _render_answer_sheet(
        self,
        doc: Document,
        grouped: list[tuple[str, str, list[dict]]],
        config: ExportConfig,
    ) -> None:
        answerable_types = {"MC", "MA", "TF", "BLANK", "SA"}
        answerable_groups: list[tuple[str, list[dict]]] = []
        for _letter, title, qs in grouped:
            items = [q for q in qs if q.get("type", "MC") in answerable_types]
            if items:
                answerable_groups.append((title, items))

        if not answerable_groups:
            return

        p = doc.add_paragraph("PHIẾU TRẢ LỜI")
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(13)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._apply_paragraph_spacing(p)

        global_counter = 0
        for title, qs in answerable_groups:
            if title:
                p = doc.add_paragraph(title)
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(12)
                self._apply_paragraph_spacing(p)

            for section_num, q in enumerate(qs, start=1):
                global_counter += 1
                section_num = global_counter if config.numbering_mode == "global" else section_num
                qtype = q.get("type", "MC")

                if qtype in ("MC", "TF"):
                    options = q.get("options", [])
                    keys = [o.get("key", "") for o in options]
                    choices = "  ".join(f"○ {k}" for k in keys)
                    p = doc.add_paragraph(
                        f"Câu {section_num}:  {choices}"
                    )
                    p.runs[0].font.size = Pt(12)
                    self._apply_paragraph_spacing(p)

                elif qtype == "MA":
                    options = q.get("options", [])
                    keys = [o.get("key", "") for o in options]
                    choices = "  ".join(f"□ {k}" for k in keys)
                    p = doc.add_paragraph(
                        f"Câu {section_num}:  {choices}"
                    )
                    p.runs[0].font.size = Pt(12)
                    self._apply_paragraph_spacing(p)

                elif qtype in ("BLANK", "SA"):
                    p = doc.add_paragraph(
                        f"Câu {section_num}:  " + "_" * 50
                    )
                    p.runs[0].font.size = Pt(12)
                    self._apply_paragraph_spacing(p)

        doc.add_paragraph("")

    # ------------------------------------------------------------------
    # Scoring rules
    # ------------------------------------------------------------------

    def _render_scoring_rules(
        self, doc: Document, questions: list[dict], config: ExportConfig | None = None
    ) -> None:
        types_present = self._types_present(questions)
        if not types_present:
            return


        p = doc.add_paragraph("QUY ĐỊNH CHẤM ĐIỂM")
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(13)
        self._apply_paragraph_spacing(p)

        rules: dict[str, str] = {
            "MC": "Mỗi câu Nhiều lựa chọn: chọn đúng 1 đáp án duy nhất mới được điểm; chọn sai hoặc bỏ trống: 0 điểm.",
            "MA": "Mỗi câu Nhiều đáp án: phải chọn đúng và đủ tất cả đáp án đúng mới được toàn bộ điểm; chọn thiếu hoặc thừa: 0 điểm.",
            "TF": "Mỗi câu Đúng/Sai: chọn đúng phương án duy nhất mới được điểm.",
            "BLANK": "Mỗi câu Điền vào chỗ trống: điền đúng (không phân biệt hoa/thường, bỏ qua khoảng trắng đầu/cuối) mới được điểm.",
            "SA": "Mỗi câu Trả lời ngắn: so khớp với đáp án mẫu (không phân biệt hoa/thường) mới được điểm.",
            "ES": "Mỗi câu CRQ - Tự luận: chấm theo thang điểm/rubric của giáo viên.",
            "PR": "Mỗi câu CRQ - Bài toán: chấm theo thang điểm/rubric của giáo viên.",
        }

        total = sum(q.get("point_value", 1.0) for q in questions)
        if config and config.essay_questions:
            total += sum(eq.get("score", 1.0) for eq in config.essay_questions)

        for qtype in ["MC", "MA", "TF", "BLANK", "SA", "ES", "PR"]:
            if qtype in types_present:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(rules[qtype])
                p.runs[0].font.size = Pt(12)
                self._apply_paragraph_spacing(p)

        if config and config.essay_questions:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run("Câu hỏi CRQ: Chấm theo thang điểm hướng dẫn chi tiết trong đáp án.")
            p.runs[0].font.size = Pt(12)
            self._apply_paragraph_spacing(p)

        p = doc.add_paragraph(f"Tổng điểm toàn bài: {total:g} điểm.")
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(12)
        self._apply_paragraph_spacing(p)
        doc.add_paragraph("")

    # ------------------------------------------------------------------
    # Answer key
    # ------------------------------------------------------------------

    def _render_answer_key(
        self,
        doc: Document,
        grouped: list[tuple[str, str, list[dict]]],
        config: ExportConfig,
    ) -> None:
        p = doc.add_paragraph("ĐÁP ÁN VÀ THANG ĐIỂM")
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(13)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._apply_paragraph_spacing(p)

        global_counter = 0
        total = 0.0

        for _letter, title, qs in grouped:
            if title:
                p = doc.add_paragraph(title)
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(12)
                self._apply_paragraph_spacing(p)

            for q in qs:
                global_counter += 1
                num = (
                    global_counter
                    if config.numbering_mode == "global"
                    else qs.index(q) + 1
                )
                qtype = q.get("type", "MC")
                points = q.get("point_value", 1.0)
                total += points

                if qtype in ("MC", "MA", "TF"):
                    correct_keys = [
                        o.get("key", "")
                        for o in q.get("options", [])
                        if o.get("is_correct", False)
                    ]
                    answer_str = ", ".join(correct_keys) if correct_keys else "(chưa có)"
                    p = doc.add_paragraph(
                        f"Câu {num}. {answer_str}  —  {points:g} điểm"
                    )
                    p.runs[0].font.size = Pt(12)
                    self._apply_paragraph_spacing(p)
                elif qtype in ("ES", "PR"):
                    p = doc.add_paragraph(
                        f"Câu {num}. {_TYPE_LABELS.get(qtype, qtype)}  —  {points:g} điểm"
                    )
                    p.runs[0].font.size = Pt(12)
                    self._apply_paragraph_spacing(p)
                    self._render_problem_template_note(doc, q, raw_latex=config.raw_latex_answer_key)
                    self._render_problem_rubric_answer_key(doc, q, raw_latex=config.raw_latex_answer_key)
                else:
                    answers = q.get("accepted_answers", [])
                    answer_str = (
                        " / ".join(
                            self._render_key_text(str(answer), raw_latex=config.raw_latex_answer_key)
                            for answer in answers
                        )
                        if answers
                        else "(chưa có)"
                    )
                    p = doc.add_paragraph(
                        f"Câu {num}. {answer_str}  —  {points:g} điểm"
                    )
                    p.runs[0].font.size = Pt(12)
                    self._apply_paragraph_spacing(p)

        doc.add_paragraph("")

        # CRQ questions in answer key
        if config.essay_questions:
            p = doc.add_paragraph("Phần CRQ")
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(12)
            self._apply_paragraph_spacing(p)
            for num, eq in enumerate(config.essay_questions, start=1):
                score = eq.get("score", 1.0)
                total += score
                p = doc.add_paragraph(
                    f"Câu {num} (CRQ): Chấm theo thang điểm  —  {score:g} điểm"
                )
                p.runs[0].font.size = Pt(12)
                self._apply_paragraph_spacing(p)

        doc.add_paragraph("")
        p = doc.add_paragraph(f"TỔNG ĐIỂM: {total:g} điểm")
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._apply_paragraph_spacing(p)

    @staticmethod
    def _is_problem_question(question: dict) -> bool:
        subtype = str(
            question.get("crq_subtype")
            or question.get("question_variant")
            or ""
        ).strip().lower()
        return subtype == "problem"

    @staticmethod
    def _is_crq_question(question: dict) -> bool:
        return str(question.get("type") or "").strip() in {"ES", "PR"}

    @staticmethod
    def _crq_subtype(question: dict) -> str:
        qtype = str(question.get("type") or "").strip()
        if qtype == "PR":
            return "problem"
        subtype = str(
            question.get("crq_subtype")
            or question.get("question_variant")
            or ""
        ).strip().lower()
        if subtype in {"essay", "problem"}:
            return subtype
        if qtype == "ES":
            return "essay"
        return ""

    def _section_title_label(self, qtype: str, questions: list[dict]) -> str:
        if qtype not in ("ES", "PR"):
            return _TYPE_LABELS.get(qtype, qtype)
        return _TYPE_LABELS.get(qtype, qtype)

    def _problem_answer_line_count(self, question: dict) -> int:
        rubric = question.get("crq_rubric") or question.get("problem_rubric") or []
        count = len(rubric)
        return count if count > 0 else 6

    def _render_numbered_answer_lines(self, doc: Document, question_number: int, question: dict) -> None:
        line_count = self._problem_answer_line_count(question)
        p = doc.add_paragraph(f"Câu {question_number}:")
        p.runs[0].font.size = Pt(12)
        table = doc.add_table(rows=line_count, cols=1)
        table.style = "Table Grid"
        table.autofit = False
        self._set_cell_width(table.cell(0, 0), 16.0)
        for row in table.rows:
            row.height = Cm(0.6)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            cell = row.cells[0]
            self._hide_cell_border(cell, "top", "left", "right")
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.left_indent = Cm(1.0)
            cell.paragraphs[0].add_run("")
        doc.add_paragraph("")

    def _render_problem_rubric_answer_key(self, doc: Document, question: dict, *, raw_latex: bool = False) -> None:
        rubric = question.get("crq_rubric") or question.get("problem_rubric") or []
        if not rubric:
            fallback = doc.add_paragraph("Chấm theo thang điểm/rubric của giáo viên.")
            fallback.runs[0].font.size = Pt(12)
            fallback.paragraph_format.left_indent = Cm(1.0)
            self._apply_paragraph_spacing(fallback)
            return

        table = doc.add_table(rows=len(rubric) + 2, cols=3)
        table.style = "Table Grid"

        widths = (2.0, 12.0, 3.0)
        for row in table.rows:
            for cell, width in zip(row.cells, widths, strict=False):
                self._set_cell_width(cell, width)

        headers = ("#", "Nội dung đáp án", "Điểm")
        for col, text in enumerate(headers):
            p = table.cell(0, col).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
            self._apply_paragraph_spacing(p)

        total_score = 0.0
        for row_index, row in enumerate(rubric, start=1):
            marker = str(row.get("marker", "") or "").strip()
            content = self._render_key_text(str(row.get("content", "") or "").strip(), raw_latex=raw_latex)
            score = float(row.get("score", 0.0) or 0.0)
            total_score += score
            values = (marker, content, self._format_score(score))
            for col, text in enumerate(values):
                p = table.cell(row_index, col).paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col != 1 else WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(text)
                run.font.size = Pt(12)
                self._apply_paragraph_spacing(p)

        total_row = len(rubric) + 1
        summary = ("TỔNG", "", self._format_score(total_score))
        for col, text in enumerate(summary):
            p = table.cell(total_row, col).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
            self._apply_paragraph_spacing(p)
        doc.add_paragraph("")

    def _render_problem_answer_sheet_line(self, doc: Document, question_number: int) -> None:
        p = doc.add_paragraph(f"Câu {question_number}:  " + "_" * 72)
        p.runs[0].font.size = Pt(12)
        self._apply_paragraph_spacing(p)

    def _render_problem_template_note(self, doc: Document, question: dict, *, raw_latex: bool = False) -> None:
        template_name = str(
            question.get("crq_template_name")
            or question.get("problem_template_name")
            or ""
        ).strip()
        template_id = question.get("crq_template_id")
        if template_id is None:
            template_id = question.get("problem_template_id")
        if not template_name:
            if template_id is None:
                return
            template_name = f"#{template_id}"
        p = doc.add_paragraph(f"Mẫu rubric: {self._render_key_text(template_name, raw_latex=raw_latex)}")
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(11)
        p.paragraph_format.left_indent = Cm(1.0)
        self._apply_paragraph_spacing(p)

    @staticmethod
    def _format_score(value: float) -> str:
        if abs(value - round(value)) < 1e-9:
            return str(int(round(value)))
        return f"{value:.3f}".rstrip("0").rstrip(".")

    def _render_question_statistics(
        self,
        doc: Document,
        questions: list[dict],
        config: ExportConfig,
    ) -> None:
        stats_questions = self._questions_for_statistics(questions, config)
        total_points = sum(float(q.get("point_value", 1.0) or 0.0) for q in stats_questions)
        sections = [
            ("Thống kê theo Chương", "Chương", self._build_statistics_rows(stats_questions, self._chapter_label)),
            ("Thống kê theo CLO", "CLO", self._build_statistics_rows(stats_questions, self._clo_label)),
            ("Thống kê theo Mức độ", "Mức độ", self._build_statistics_rows(stats_questions, self._difficulty_label)),
            ("Thống kê theo Loại câu hỏi", "Loại", self._build_statistics_rows(stats_questions, self._type_label)),
        ]

        for index, (title, group_header, rows) in enumerate(sections):
            p = doc.add_paragraph(title)
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(13)
            self._render_statistics_table(doc, group_header, rows, total_points)
            if index < len(sections) - 1:
                doc.add_paragraph("")

    def _render_statistics_table(
        self,
        doc: Document,
        group_header: str,
        rows: list[tuple[str, int, float]],
        total_points: float,
    ) -> None:
        table = doc.add_table(rows=len(rows) + 2, cols=4)
        table.style = "Table Grid"

        widths = (6.0, 3.0, 3.0, 4.0)
        for row in table.rows:
            for cell, width in zip(row.cells, widths, strict=False):
                self._set_cell_width(cell, width)

        headers = (group_header, "Số câu", "Điểm", "Tỷ lệ điểm")
        for col, text in enumerate(headers):
            p = table.cell(0, col).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)

        total_count = 0
        total_score = 0.0
        for row_index, (label, count, score) in enumerate(rows, start=1):
            total_count += count
            total_score += score
            values = (
                label,
                str(count),
                f"{score:g}",
                self._ratio_text(score, total_points),
            )
            for col, text in enumerate(values):
                p = table.cell(row_index, col).paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.font.size = Pt(12)

        total_row = len(rows) + 1
        summary = (
            "Tổng",
            str(total_count),
            f"{total_score:g}",
            self._ratio_text(total_score, total_points),
        )
        for col, text in enumerate(summary):
            p = table.cell(total_row, col).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)

    def _build_statistics_rows(
        self,
        questions: list[dict],
        label_getter,
    ) -> list[tuple[str, int, float]]:
        stats: dict[str, tuple[int, float]] = {}
        for question in questions:
            label = label_getter(question)
            count, score = stats.get(label, (0, 0.0))
            stats[label] = (count + 1, score + float(question.get("point_value", 1.0) or 0.0))
        return [(label, count, score) for label, (count, score) in stats.items()]

    def _questions_for_statistics(self, questions: list[dict], config: ExportConfig) -> list[dict]:
        stats_questions = [dict(question) for question in questions]
        for essay in config.essay_questions:
            stats_questions.append(
                {
                    "type": "ES",
                    "content": "",
                    "point_value": essay.get("score", 1.0),
                    "difficulty": "",
                    "learning_outcome_code": "",
                    "category": "",
                }
            )
        return stats_questions

    @staticmethod
    def _ratio_text(score: float, total_points: float) -> str:
        if total_points <= 0:
            return "0.0%"
        return f"{(score / total_points) * 100:.1f}%"

    @staticmethod
    def _chapter_label(question: dict) -> str:
        value = str(question.get("category") or "").strip()
        return value or "(Chưa gắn chương)"

    @staticmethod
    def _clo_label(question: dict) -> str:
        value = str(question.get("learning_outcome_code") or "").strip()
        return value or "(Chưa gắn CLO)"

    @staticmethod
    def _difficulty_label(question: dict) -> str:
        value = str(question.get("difficulty") or "").strip()
        return _DIFFICULTY_LABELS.get(value, value or "(Chưa gắn mức độ)")

    @staticmethod
    def _type_label(question: dict) -> str:
        value = str(question.get("type") or "").strip()
        return _TYPE_STATS_LABELS.get(value, value or "(Chưa gắn loại)")

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _types_present(self, questions: list[dict]) -> set[str]:
        return {q.get("type", "MC") for q in questions}


# ---------------------------------------------------------------------------
# Convenience function
# ---------------------------------------------------------------------------

def build_output_path(title: str, exports_dir: Path) -> Path:
    """Generate a timestamped output path inside *exports_dir*.

    Example: ``exports/MyExam_2026-04-21_143022.docx``
    """
    exports_dir.mkdir(parents=True, exist_ok=True)
    safe_title = re.sub(r'[\\/:*?"<>|]', "_", title).strip() or "exam"
    timestamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    return exports_dir / f"{safe_title}_{timestamp}.docx"
